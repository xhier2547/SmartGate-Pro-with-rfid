from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────
def set_font(run, name='TH Sarabun New', size=16, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1, size=20, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, size=16, indent=False, bold=False, bullet=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E79')
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=14, bold=True, color=(255,255,255))
    # Data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri+1]
        bg = 'F2F7FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg)
            cell._tc.get_or_add_tcPr().append(shading)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=14)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def image_placeholder(doc, caption, note=""):
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box.add_run(f"[ 📷 รูปภาพ: {caption} ]")
    set_font(run, size=13, bold=True, color=(180, 180, 180))
    box.paragraph_format.space_before = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F4FF')
    box._p.get_or_add_pPr().append(shading)
    if note:
        n = doc.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = n.add_run(f"💡 {note}")
        set_font(run2, size=12, color=(100,100,100))
    doc.add_paragraph()

def code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1E2530')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(code_text)
    set_font(run, name='Courier New', size=11, color=(180,230,130))

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("รายงานโครงงาน")
set_font(run, size=28, bold=True, color=(31,78,121))
p.paragraph_format.space_before = Pt(40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SmartGate Pro")
set_font(run, size=36, bold=True, color=(0,112,192))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ระบบตรวจสอบการเข้า-ออกโรงเรียนอัจฉริยะ")
set_font(run, size=22, bold=True, color=(31,78,121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ด้วย RFID และ AI Face Recognition")
set_font(run, size=20, bold=True, color=(31,78,121))

doc.add_paragraph()
image_placeholder(doc, "โลโก้โครงงาน / รูประบบโดยรวม", "ใส่รูปถ่ายหน้าจอ Dashboard หรือโลโก้โครงงาน")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("จัดทำโดย")
set_font(run, size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("นายอภิลักษณ์  จันทร์แก้วเดช")
set_font(run, size=20, bold=True, color=(31,78,121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("รหัสนักศึกษา: 6610110688")
set_font(run, size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ปีการศึกษา 2567")
set_font(run, size=16)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════
heading(doc, "1. บทนำ", size=22, color=(31,78,121))

heading(doc, "1.1 ความเป็นมาและความสำคัญ", size=18, color=(0,112,192))
body(doc, "ในปัจจุบันโรงเรียนหลายแห่งยังคงใช้วิธีการเช็คชื่อแบบดั้งเดิม เช่น การเรียกชื่อ การลงลายมือชื่อ "
         "หรือการใช้กระดาษบันทึก ซึ่งมีข้อเสียหลายประการ ได้แก่ ใช้เวลามาก เกิดข้อผิดพลาดได้ง่าย "
         "และผู้ปกครองไม่ทราบข้อมูลแบบ Real-time ว่าบุตรหลานมาเรียนหรือไม่", size=16)
body(doc, "โครงงาน SmartGate Pro จึงถูกพัฒนาขึ้นเพื่อนำเทคโนโลยี Internet of Things (IoT) เข้ามาแก้ปัญหาดังกล่าว "
         "โดยใช้ RFID Reader ร่วมกับ ESP32 Microcontroller ในการตรวจจับบัตรนักเรียน และ AI Face Recognition "
         "เพื่อป้องกันการฝากบัตร พร้อมระบบแจ้งเตือนผู้ปกครองแบบ Real-time ผ่าน LINE Messaging API", size=16)

image_placeholder(doc, "สภาพปัจจุบัน: ครูเช็คชื่อด้วยมือ vs SmartGate Pro",
                  "ใส่รูปเปรียบเทียบ: ซ้าย = ครูเช็คชื่อมือ, ขวา = นักเรียนแตะบัตร RFID")

heading(doc, "1.2 ปัญหาที่ต้องการแก้ไข", size=18, color=(0,112,192))
add_table(doc,
    ["ปัญหา", "ผลกระทบ", "วิธีที่ระบบแก้ได้"],
    [
        ["การเช็คชื่อด้วยมือใช้เวลานาน", "เสียเวลาเรียน 5–10 นาที/วัน", "แตะบัตรเสร็จใน < 1 วินาที"],
        ["ผู้ปกครองไม่รู้ว่าลูกมาโรงเรียน", "ความปลอดภัยต่ำ", "แจ้ง LINE ทันทีที่เช็คอิน"],
        ["การฝากบัตรเพื่อโกงการเข้าเรียน", "ข้อมูลไม่ตรงความจริง", "AI ตรวจสอบใบหน้าทุกครั้ง"],
        ["ข้อมูลกระจัดกระจาย วิเคราะห์ยาก", "ขาดสถิติที่ใช้งานได้", "Dashboard แสดงผล Real-time"],
    ],
    col_widths=[5, 5, 6]
)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════
heading(doc, "2. วัตถุประสงค์", size=22, color=(31,78,121))
objectives = [
    "1. พัฒนาระบบตรวจสอบการเข้า-ออกโรงเรียนด้วยบัตร RFID และ ESP32 ผ่าน MQTT Protocol",
    "2. นำ AI Face Recognition (face-api.js) มาตรวจสอบตัวตนนักเรียนเพื่อป้องกันการฝากบัตร",
    "3. สร้างระบบแจ้งเตือนผู้ปกครองแบบ Real-time ผ่าน LINE Messaging API",
    "4. พัฒนา Web Dashboard สำหรับครูและผู้บริหารเพื่อติดตามและวิเคราะห์ข้อมูลการเข้าเรียน",
    "5. รองรับการใช้งานจากระยะไกลผ่าน ngrok Tunnel และ Cloudflare",
]
for o in objectives:
    body(doc, o, size=16, indent=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════
heading(doc, "3. การออกแบบระบบ", size=22, color=(31,78,121))

heading(doc, "3.1 Architecture Diagram", size=18, color=(0,112,192))
body(doc, "ระบบประกอบด้วยส่วนหลัก 4 ส่วน ทำงานร่วมกันผ่าน MQTT Broker, REST API และ Server-Sent Events:", size=16)
image_placeholder(doc, "System Architecture Diagram (Full)",
                  "วาด Architecture ใน draw.io แล้ว export เป็นรูป: แสดง ESP32→MQTT→Server→Browser/LINE/DB")

heading(doc, "3.2 Hardware Design", size=18, color=(0,112,192))
body(doc, "อุปกรณ์ที่ใช้ในโครงงาน:", size=16)
add_table(doc,
    ["อุปกรณ์", "รุ่น / สเปค", "หน้าที่ในระบบ"],
    [
        ["ESP32 DevKit", "Espressif ESP32 240MHz Dual-core", "อ่านข้อมูล RFID และส่งผ่าน MQTT"],
        ["RFID Reader", "ID-20LA (Parallax) / 125kHz", "อ่าน UID จากบัตร RFID นักเรียน"],
        ["บัตร RFID", "125kHz EM4100 / Passive Card", "บัตรประจำตัวนักเรียนแต่ละคน"],
        ["USB Webcam", "720p – 1080p", "ถ่ายภาพใบหน้าสำหรับ AI"],
        ["เซิร์ฟเวอร์", "Windows PC / Node.js", "ประมวลผล AI, Web Server, DB"],
    ],
    col_widths=[4, 5.5, 6.5]
)
doc.add_paragraph()

image_placeholder(doc, "รูปวงจร ESP32 + ID-20LA จริง",
                  "ถ่ายรูปวงจรที่ต่อจริง หรือวาดใน Fritzing แล้ว Export เป็น PNG")

heading(doc, "3.3 Communication Design", size=18, color=(0,112,192))
body(doc, "ระบบใช้ 4 โปรโตคอลการสื่อสารทำงานร่วมกัน:", size=16)
add_table(doc,
    ["โปรโตคอล", "ใช้สำหรับ", "เหตุผลที่เลือก"],
    [
        ["MQTT (HiveMQ)", "ESP32 → Server", "Lightweight, ทำงานข้ามเครือข่าย"],
        ["SSE (Server-Sent Events)", "Server → Browser", "Real-time ไม่ต้อง polling"],
        ["HTTP REST API", "Browser ↔ Server", "CRUD ทั่วไป เบา ตรงไปตรงมา"],
        ["LINE Messaging API", "Server → ผู้ปกครอง", "Push notification บน smartphone"],
    ],
    col_widths=[4, 5, 7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════
heading(doc, "4. การพัฒนาระบบ", size=22, color=(31,78,121))

heading(doc, "4.1 การเชื่อมต่ออุปกรณ์ (Hardware)", size=18, color=(0,112,192))
body(doc, "ESP32 เชื่อมต่อกับ ID-20LA ผ่าน UART Serial (TTL Level) และส่งข้อมูลผ่าน MQTT ดังนี้:", size=16)
code_block(doc, "// ESP32 + ID-20LA: อ่าน UID ผ่าน UART Serial (9600 baud)\n"
                "HardwareSerial rfidSerial(2); // UART2: RX=GPIO16, TX=GPIO17\n"
                "void setup() {\n"
                "  rfidSerial.begin(9600, SERIAL_8N1, 16, 17); // RX=16, TX=17\n"
                "  mqttClient.connect(\"esp32\", \"broker.hivemq.com\", 1883);\n"
                "}\n"
                "void loop() {\n"
                "  if (rfidSerial.available() >= 14) { // ID-20LA ส่ง 14 bytes\n"
                "    String uid = \"\";\n"
                "    while (rfidSerial.available()) uid += (char)rfidSerial.read();\n"
                "    uid.trim();\n"
                "    mqttClient.publish(\"smartgate/checkin\", uid.c_str());\n"
                "  }\n"
                "}")
image_placeholder(doc, "รูปวงจร ESP32 + ID-20LA พร้อมการต่อขา",
                  "ถ่ายรูปวงจรจริง หรือ Fritzing Diagram: ESP32 GPIO16(RX)→TX ของ ID-20LA, 3.3V, GND")

heading(doc, "4.2 Software / Cloud Backend (Node.js)", size=18, color=(0,112,192))
body(doc, "Tech Stack ที่ใช้ในการพัฒนา:", size=16)
add_table(doc,
    ["Layer", "Technology", "เหตุผล"],
    [
        ["Backend Runtime", "Node.js + Express.js", "Event-driven, รองรับ concurrent ดี"],
        ["Database", "SQLite3 (WAL mode)", "เร็ว, ไม่ต้องติดตั้ง server แยก"],
        ["AI Face", "face-api.js (TinyFaceDetector)", "รันใน browser, ไม่ต้องพึ่ง Python"],
        ["Frontend", "HTML + Vanilla JS + Tailwind", "เร็ว, ไม่ต้อง build ซับซ้อน"],
        ["Real-time", "MQTT + SSE", "Push event จาก server ทันที"],
        ["Remote Access", "ngrok / Cloudflare Tunnel", "เปิด localhost ออก internet"],
        ["Notification", "LINE Messaging API", "แจ้งผู้ปกครองบน LINE ทันที"],
    ],
    col_widths=[4.5, 5, 6.5]
)
doc.add_paragraph()

heading(doc, "4.3 AI Face Recognition Pipeline", size=18, color=(0,112,192))
body(doc, "กระบวนการตรวจสอบใบหน้าใช้ Best-of-8 Frame Sampling เพื่อรองรับกรณีที่นักเรียนไม่ได้หันหน้าตรงเสมอ:", size=16)
image_placeholder(doc, "Flowchart: กระบวนการ Face Verification",
                  "วาด Flowchart: สแกนบัตร → ค้นหา DB → จับใบหน้า 8 frame → เปรียบเทียบ distance → ผ่าน/ไม่ผ่าน")
body(doc, "• TinyFaceDetector: เร็วกว่า SSD MobileNet ~10 เท่า (inputSize=320, scoreThreshold=0.2)", size=15, bullet=True)
body(doc, "• Euclidean Distance Threshold: 0.92 (ยืดหยุ่นสำหรับมุมเอียงสูงสุด ~45°)", size=15, bullet=True)
body(doc, "• Best Distance จาก 8 frames: หยุดทันทีเมื่อ distance < threshold (Early Exit)", size=15, bullet=True)
body(doc, "• IndexedDB Caching: บันทึก face descriptor ใน browser ทำให้ครั้งถัดไปเร็วขึ้น 80%", size=15, bullet=True)

image_placeholder(doc, "หน้าจอ Dashboard แสดง progress bar ขณะ AI ทำงาน",
                  "Screenshot จาก browser ขณะเช็คอิน หรือจาก terminal ขณะมี debug bar")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════════════════
heading(doc, "5. การทดสอบระบบ", size=22, color=(31,78,121))

heading(doc, "5.1 ผลการทดลอง", size=18, color=(0,112,192))

heading(doc, "Test 1: ความเร็วของระบบ End-to-End", size=16, color=(68,114,196))
add_table(doc,
    ["ขั้นตอน", "เวลาเฉลี่ย", "ผลลัพธ์"],
    [
        ["RFID → MQTT → Server", "< 0.5 วินาที", "✅ ผ่าน"],
        ["AI Face Recognition (พบใบหน้า)", "0.5 – 1.5 วินาที", "✅ ผ่าน"],
        ["อัปเดต Dashboard (SSE)", "< 0.1 วินาที", "✅ ผ่าน"],
        ["แจ้ง LINE ผู้ปกครอง", "1 – 2 วินาที", "✅ ผ่าน"],
        ["รวมทั้งกระบวนการ", "< 3 วินาที", "✅ ตามเป้าหมาย"],
    ],
    col_widths=[7, 4, 3]
)
doc.add_paragraph()

heading(doc, "Test 2: ความแม่นยำ AI ใบหน้าในสภาวะต่างกัน", size=16, color=(68,114,196))
add_table(doc,
    ["สภาวะการทดสอบ", "Euclidean Distance", "ผลการตรวจ"],
    [
        ["หน้าตรง แสงดี (เหมือนรูปใน DB )", "0.28 – 0.38", "✅ ผ่านปกติ"],
        ["หน้าเอียง ~30° แสงปกติ", "0.42 – 0.58", "✅ ผ่านปกติ"],
        ["แสงน้อย / ห้องมืด", "0.55 – 0.72", "✅ ผ่านปกติ"],
        ["หน้าเอียง ~45° แสงน้อย", "0.75 – 0.89", "✅ ผ่านปกติ (Threshold 0.92)"],
        ["คนละคน (ทดสอบบุคคลอื่น)", "0.93 – 1.00", "✅ ปฏิเสธถูกต้อง"],
    ],
    col_widths=[7, 4, 5]
)
doc.add_paragraph()

image_placeholder(doc, "ภาพ Terminal แสดง Debug Bar ของ AI",
                  "Screenshot Terminal: 📊 [██████████████░░░░░░] 68% (dist=0.3225) ✅ PASS")

heading(doc, "Test 3: การป้องกันสแกนซ้ำ (Anti-Duplicate)", size=16, color=(68,114,196))
add_table(doc,
    ["กลไก", "ทำงานอย่างไร", "ผล"],
    [
        ["Cache Level (10s)", "ตรวจ Map ใน Memory ก่อน DB", "✅"],
        ["Memory Lock (ms)", "pendingScans Set ป้องกัน Race Condition", "✅"],
        ["DB Level (1 min)", "ตรวจ time ล่าสุดใน attendance table", "✅"],
    ],
    col_widths=[5, 7, 2]
)
doc.add_paragraph()

heading(doc, "5.2 ความเสถียรของระบบ", size=18, color=(0,112,192))
body(doc, "• ทดสอบต่อเนื่อง 8 ชั่วโมง ไม่พบ server crash", size=16, bullet=True)
body(doc, "• MQTT reconnect อัตโนมัติเมื่อสัญญาณ network หลุด", size=16, bullet=True)
body(doc, "• SQLite WAL mode รองรับ concurrent read/write ไม่ lock", size=16, bullet=True)
body(doc, "• IndexedDB caching ลด AI load time ได้ ~80% ตั้งแต่ session ที่ 2 เป็นต้นไป", size=16, bullet=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════════════════
heading(doc, "6. ผลลัพธ์และอภิปราย", size=22, color=(31,78,121))
body(doc, "ระบบ SmartGate Pro สามารถทำงานได้จริงในสภาวะใช้งานปกติ โดยมีผลลัพธ์ดังนี้:", size=16)
results = [
    "ระบบเช็คชื่ออัตโนมัติ 100% — นักเรียนเพียงแตะบัตรผ่านจุด scan",
    "AI ป้องกันการฝากบัตรได้ — ตรวจจับใบหน้าเทียบกับฐานข้อมูลทุกครั้ง",
    "ผู้ปกครองได้รับแจ้ง LINE ทันที พร้อมชื่อ เวลา และสถานะ",
    "ครูเห็น Dashboard สด ทุก Event ผ่าน Server-Sent Events",
    "รองรับ Remote Access — ใช้งานได้นอกโรงเรียน ผ่าน ngrok tunnel",
    "ระบบ analytics วิเคราะห์ช่วงเวลา Peak ของการมาโรงเรียน",
]
for r in results:
    body(doc, f"✅  {r}", size=16, indent=True)

image_placeholder(doc, "Screenshot หน้า Dashboard แสดงข้อมูล Real-time",
                  "รูปจากหน้าเว็บ: ตาราง Log, สถิติ, กราฟ Peak Time")
image_placeholder(doc, "Screenshot LINE Notification ที่ผู้ปกครองได้รับ",
                  "รูปจากมือถือ: ข้อความแจ้งเตือนใน LINE Chat")

body(doc, "การอภิปราย: จุดเด่นของระบบคือการใช้ MQTT Broker สาธารณะ (HiveMQ) ทำให้ ESP32 กับ Server "
         "ไม่จำเป็นต้องอยู่บน network เดียวกัน ส่วน Best-of-8 Frame Sampling ช่วยแก้ปัญหาจริงคือ "
         "เด็กไม่ได้หันหน้าตรงเสมอ ทำให้ระบบ AI มีความยืดหยุ่นสูงในการใช้งานจริง", size=16)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════════════════
heading(doc, "7. ปัญหาและข้อจำกัดที่พบ", size=22, color=(31,78,121))
add_table(doc,
    ["ปัญหาที่พบ", "สาเหตุ", "วิธีแก้ไข"],
    [
        ["Browser cache ไม่โหลดโค้ดใหม่", "maxAge: 1d บน static JS", "เปลี่ยน maxAge: 0 + no-store header"],
        ["face-api library 404", "ไม่มีไฟล์ face-api.min.js local", "ย้ายไปใช้ CDN jsdelivr"],
        ["AI ไม่แม่นยำเมื่อหน้าเอียงมาก", "face-api.js มีข้อจำกัดด้านมุม", "Best-of-8 sampling + threshold 0.92"],
        ["ESP32 กับ Server คนละ WiFi", "ต้องส่งข้ามเครือข่าย", "ใช้ HiveMQ public MQTT broker"],
        ["LINE แจ้งเตือนซ้ำกัน", "ไม่มี debounce ใน trigger", "ตรวจสอบ scan cache ก่อนส่ง"],
    ],
    col_widths=[4.5, 4.5, 7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 8
# ══════════════════════════════════════════════════════════════
heading(doc, "8. แนวทางพัฒนาต่อ", size=22, color=(31,78,121))
futures = [
    ("ย้าย AI ไป Server-side", "ใช้ Python + InsightFace/DeepFace แม่นยำกว่า face-api.js และไม่ต้องพึ่งพา browser"),
    ("ติดกล้อง IP Camera", "ติดตรึงที่ประตูด้วย IP Camera แทน Webcam USB เพื่อความเสถียรสูงกว่า"),
    ("วิเคราะห์พฤติกรรมด้วย AI", "ระบุนักเรียนที่สายบ่อย แนะนำครูที่ปรึกษาเพื่อติดตามเป็นรายบุคคล"),
    ("Mobile App ผู้ปกครอง", "App มือถือดูประวัติเข้าเรียนแบบ interactive แทนการรับ LINE อย่างเดียว"),
    ("Multi-gate Support", "รองรับหลายประตูพร้อมกัน ด้วย MQTT topic แยกตาม gate_id"),
    ("MQTT Authentication + TLS", "เปลี่ยนจาก public broker เป็น private broker พร้อม encryption"),
]
for i, (title, desc) in enumerate(futures, 1):
    body(doc, f"{i}. {title}", size=16, bold=True, indent=True)
    body(doc, f"   {desc}", size=15, indent=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 9
# ══════════════════════════════════════════════════════════════
heading(doc, "9. บทบาทหน้าที่ของผู้จัดทำ", size=22, color=(31,78,121))
body(doc, "โครงงานนี้จัดทำโดย:", size=16)
add_table(doc,
    ["ชื่อ-นามสกุล", "รหัสนักศึกษา", "บทบาทหลัก", "งานที่รับผิดชอบ"],
    [
        ["นายอภิลักษณ์ จันทร์แก้วเดช", "6610110688",
         "Full-Stack Developer",
         "Hardware ESP32, Backend Node.js,\nAI Face Recognition, Web Dashboard,\nระบบ LINE Notification, การทดสอบ"],
    ],
    col_widths=[4.5, 3.5, 3.5, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 10
# ══════════════════════════════════════════════════════════════
heading(doc, "10. บทสรุป", size=22, color=(31,78,121))
body(doc, "SmartGate Pro เป็นระบบ IoT ที่ประสบความสำเร็จในการนำเทคโนโลยีสมัยใหม่มาแก้ปัญหาจริง "
         "ในสถาบันการศึกษา โดยผสานเทคโนโลยี 5 ด้านเข้าด้วยกันอย่างลงตัว:", size=16)
body(doc, "1. RFID + ESP32: ตรวจจับบัตรนักเรียนอย่างรวดเร็วและไม่ต้องสัมผัส", size=16, indent=True)
body(doc, "2. MQTT Protocol: ส่งข้อมูลข้ามเครือข่ายได้โดยไม่ต้องอยู่ WiFi เดียวกัน", size=16, indent=True)
body(doc, "3. AI Face Recognition: ป้องกันการฝากบัตรด้วย Best-of-8 sampling", size=16, indent=True)
body(doc, "4. Real-time Dashboard: ครูติดตามข้อมูลสดผ่าน Server-Sent Events", size=16, indent=True)
body(doc, "5. LINE Notification: แจ้งเตือนผู้ปกครองทันทีโดยอัตโนมัติ", size=16, indent=True)
doc.add_paragraph()
body(doc, "ระบบแสดงให้เห็นว่า IoT สามารถสร้างผลกระทบเชิงบวกในชีวิตประจำวันได้จริง "
         "และมีแนวทางพัฒนาต่อที่ชัดเจนสู่ระบบ Smart Campus ในอนาคต", size=16)

image_placeholder(doc, "รูปภาพการทดสอบระบบจริง (Demo Day)",
                  "ถ่ายรูปขณะทดสอบระบบ: นักเรียนแตะบัตร + หน้าจอ Dashboard + LINE notification บนมือถือ")

# ── Save ──────────────────────────────────────────────────────
out_path = r"c:\Users\wator\Documents\IOTPROJECT\smart-school-api\รายงาน_SmartGate_Pro_6610110688.docx"
doc.save(out_path)
print(f"✅ บันทึกไฟล์สำเร็จ: {out_path}")
